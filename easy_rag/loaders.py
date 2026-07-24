"""Load documents of several file types from a path into Document objects.

Supported out of the box: .txt, .md, .markdown, .rst, .csv (core, no extra
dependencies). .pdf, .docx, and images (.png/.jpg/.jpeg/.bmp/.tiff) need one
extra each -- see the ImportError messages below for the exact install
command; the base install works without any of them.
"""

import csv
import os
import tempfile
from dataclasses import dataclass, field


@dataclass
class Document:
    text: str
    source: str
    metadata: dict = field(default_factory=dict)


TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".rst"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tiff"}


def _load_text_file(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _load_csv_file(path):
    with open(path, "r", encoding="utf-8", errors="ignore", newline="") as f:
        reader = csv.reader(f)
        rows = list(reader)
    if not rows:
        return ""
    header, data_rows = rows[0], rows[1:]
    lines = []
    for row in data_rows:
        pairs = ", ".join(f"{h}: {v}" for h, v in zip(header, row))
        lines.append(pairs)
    # Each row is its own paragraph (chunking.py splits on blank lines) so
    # the chunker breaks between rows instead of hard-wrapping mid-row --
    # a single "\n" join collapses the whole CSV into one giant paragraph,
    # which for punctuation-free "key: value" rows falls through the
    # sentence splitter untouched and gets cut by raw character count,
    # slicing words in half across chunk boundaries.
    return "\n\n".join(lines)


def _load_pdf_file(path):
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ImportError(
            "Reading PDF files requires pypdf. Install it with: pip install easy-rag[pdf]"
        ) from e
    reader = PdfReader(path)
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def _read_opendataloader_text(paths, out_dir):
    results = {}
    for p in paths:
        stem = os.path.splitext(os.path.basename(p))[0]
        text_path = os.path.join(out_dir, stem + ".txt")
        if os.path.exists(text_path):
            with open(text_path, "r", encoding="utf-8", errors="ignore") as f:
                results[p] = f.read()
    return results


def _convert_one_opendataloader(opendataloader_pdf, path, quiet):
    with tempfile.TemporaryDirectory() as out_dir:
        try:
            opendataloader_pdf.convert(input_path=[path], output_dir=out_dir, format="text", quiet=quiet)
        except FileNotFoundError:
            # Java itself isn't installed/on PATH -- a setup problem
            # affecting every file identically, not this one file being
            # bad. Must not be swallowed as a per-file skip: silently
            # treating "Java is missing" as "every PDF happened to be
            # corrupt" would drop all PDF content with no diagnostic at
            # all pointing at the real, fixable cause.
            raise
        except Exception:
            return None
        return _read_opendataloader_text([path], out_dir).get(path)


def _convert_pdfs_opendataloader(pdf_paths, quiet=True):
    """Batch-convert PDFs via opendataloader-pdf, extracting text with
    reading order and structure preserved -- generally higher-quality than
    pypdf's flat text dump, at the cost of requiring Java 11+ on the system
    (it wraps a JVM-based parser under the hood; each call spawns a JVM, so
    batching many files into one call is much faster than one call per file).

    Returns {path: extracted_text} for every path that converted
    successfully; a path that failed to convert (corrupt, encrypted,
    unreadable, ...) is simply absent, matching every other loader's
    per-file resilience.

    Two things this guards against, found while testing this integration
    against the real tool rather than just its docs:
      - One invalid PDF in a batch call makes the *entire* call fail with
        no output for any file, not just the bad one -- falls back to
        converting one file at a time so only the genuinely bad ones are
        skipped.
      - Every file in one batch call writes its output into the same
        output_dir, named only after the input file's basename with no
        directory structure preserved -- two different PDFs that happen to
        share a basename (e.g. "report.pdf" in two different folders)
        would silently overwrite each other's output if batched together.
        Basename-colliding files are therefore always converted one at a
        time, each into its own isolated temp directory.
    """
    try:
        import opendataloader_pdf
    except ImportError as e:
        raise ImportError(
            "The 'opendataloader' PDF backend requires the opendataloader-pdf "
            "package (and Java 11+ installed on your system). "
            "Install the package with: pip install easy-rag[opendataloader]"
        ) from e

    by_basename = {}
    for p in pdf_paths:
        by_basename.setdefault(os.path.basename(p), []).append(p)
    batchable = [group[0] for group in by_basename.values() if len(group) == 1]
    colliding = [p for group in by_basename.values() if len(group) > 1 for p in group]

    try:
        results = {}
        if batchable:
            with tempfile.TemporaryDirectory() as out_dir:
                try:
                    opendataloader_pdf.convert(input_path=list(batchable), output_dir=out_dir, format="text", quiet=quiet)
                    results.update(_read_opendataloader_text(batchable, out_dir))
                except FileNotFoundError:
                    raise  # see below -- a setup problem, not a bad file
                except Exception:
                    for p in batchable:
                        text = _convert_one_opendataloader(opendataloader_pdf, p, quiet)
                        if text is not None:
                            results[p] = text

        for p in colliding:
            text = _convert_one_opendataloader(opendataloader_pdf, p, quiet)
            if text is not None:
                results[p] = text
    except FileNotFoundError as e:
        raise RuntimeError(
            "The 'opendataloader' PDF backend requires Java 11+ installed and on "
            "your system PATH (opendataloader-pdf wraps a JVM-based parser, not a "
            "pure-Python one). Install Java, or use pdf_backend='pypdf' instead."
        ) from e

    return results


def _load_docx_file(path):
    try:
        from docx import Document as DocxDocument
    except ImportError as e:
        raise ImportError(
            "Reading .docx files requires python-docx. Install it with: pip install easy-rag[docx]"
        ) from e
    doc = DocxDocument(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text for cell in row.cells))
    return "\n\n".join(paragraphs)


def _load_image_file(path):
    try:
        import pytesseract
        from PIL import Image
    except ImportError as e:
        raise ImportError(
            "Reading images requires pytesseract and Pillow, plus the Tesseract "
            "OCR engine installed on your system (https://github.com/tesseract-ocr/tesseract). "
            "Install the Python packages with: pip install easy-rag[ocr]"
        ) from e
    return pytesseract.image_to_string(Image.open(path))


_LOADERS = {
    ".csv": _load_csv_file,
    ".pdf": _load_pdf_file,
    ".docx": _load_docx_file,
}
for _ext in IMAGE_EXTENSIONS:
    _LOADERS[_ext] = _load_image_file


def load_documents(path, pdf_backend="pypdf"):
    """Load every supported file under `path` (a single file or a directory,
    searched recursively) into a list of Document objects. Unsupported file
    types are silently skipped.

    pdf_backend selects how .pdf files are read:
      - "pypdf" (default): lightweight, pure Python, no extra system
        dependency beyond `pip install easy-rag[pdf]`.
      - "opendataloader": higher-quality extraction with reading order and
        document structure preserved, via the opendataloader-pdf project
        (https://github.com/opendataloader-project/opendataloader-pdf).
        Requires `pip install easy-rag[opendataloader]` *and* Java 11+
        installed on the system -- opendataloader-pdf wraps a JVM-based
        parser, not a pure-Python one.
    """
    if pdf_backend not in ("pypdf", "opendataloader"):
        raise ValueError(f"Unknown pdf_backend {pdf_backend!r}. Choose 'pypdf' or 'opendataloader'.")

    paths = []
    if os.path.isfile(path):
        paths = [path]
    elif os.path.isdir(path):
        for root, _dirs, files in os.walk(path):
            for name in files:
                paths.append(os.path.join(root, name))
    else:
        raise FileNotFoundError(f"No such file or directory: {path}")

    opendataloader_texts = {}
    if pdf_backend == "opendataloader":
        pdf_paths = [p for p in paths if os.path.splitext(p)[1].lower() == ".pdf"]
        if pdf_paths:
            opendataloader_texts = _convert_pdfs_opendataloader(pdf_paths)

    documents = []
    for p in sorted(paths):
        ext = os.path.splitext(p)[1].lower()
        if ext == ".pdf" and pdf_backend == "opendataloader":
            text = opendataloader_texts.get(p, "")
            if text.strip():
                documents.append(Document(text=text, source=p))
            continue
        if ext in TEXT_EXTENSIONS:
            loader = _load_text_file
        elif ext in _LOADERS:
            loader = _LOADERS[ext]
        else:
            continue
        try:
            text = loader(p)
        except ImportError:
            # A missing optional dependency is a setup problem the user
            # needs to fix (the error message tells them the exact `pip
            # install` command) -- never swallow it silently.
            raise
        except Exception:
            # Any other failure (corrupt file, unreadable/locked file, OCR
            # failure, ...) means this one file is unreadable, not that the
            # whole batch should be discarded -- skip it and keep going, so
            # one bad file doesn't lose every other valid file in the folder.
            continue
        if text.strip():
            documents.append(Document(text=text, source=p))
    return documents
